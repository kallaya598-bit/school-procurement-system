import io
import json
import os
import re
import uuid
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import quote, unquote, urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE_DIR = Path(__file__).resolve().parent
GENERATED_DIR = BASE_DIR / "generated"
GENERATED_DIR.mkdir(exist_ok=True)
FONT_NAME = "TH SarabunPSK"
GARUDA_PATH = BASE_DIR / "images.jpg"

THAI_MONTHS = ["","มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน","กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]


def clean_text(value, default=""):
    if value is None:
        return default
    return str(value).strip()


def money(value):
    try:
        return float(str(value).replace(",", "") or 0)
    except ValueError:
        return 0.0


def fmt_money(value):
    return f"{money(value):,.2f}"


def thai_number_text(number):
    number = int(number)
    digits = ["ศูนย์","หนึ่ง","สอง","สาม","สี่","ห้า","หก","เจ็ด","แปด","เก้า"]
    units = ["","สิบ","ร้อย","พัน","หมื่น","แสน","ล้าน"]
    if number == 0:
        return digits[0]
    def read_group(n):
        text = ""
        s = str(n)
        length = len(s)
        for i, ch in enumerate(s):
            d = int(ch)
            pos = length - i - 1
            if d == 0:
                continue
            if pos == 0 and d == 1 and length > 1:
                text += "เอ็ด"
            elif pos == 1 and d == 1:
                text += "สิบ"
            elif pos == 1 and d == 2:
                text += "ยี่สิบ"
            else:
                text += digits[d] + units[pos]
        return text
    parts = []
    while number:
        parts.append(number % 1_000_000)
        number //= 1_000_000
    out = ""
    for idx in reversed(range(len(parts))):
        if parts[idx]:
            out += read_group(parts[idx])
            if idx:
                out += "ล้าน"
    return out


def baht_text(value):
    amount = round(money(value) + 0.00001, 2)
    baht = int(amount)
    satang = int(round((amount - baht) * 100))
    text = thai_number_text(baht) + "บาท"
    if satang:
        text += thai_number_text(satang) + "สตางค์"
    else:
        text += "ถ้วน"
    return text


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_cm):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table, widths_cm):
    """บังคับ layout แบบ fixed + กำหนด tblGrid ให้ LibreOffice/Word เคารพความกว้างคอลัมน์
    (กัน LibreOffice บีบคอลัมน์จนข้อความตัดบรรทัดเป็นหลายบรรทัด ทำให้แถวสูงเกิน)"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    # 1) tblLayout = fixed
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    # 2) tblGrid: กำหนดความกว้างแต่ละคอลัมน์ให้ชัดเจน
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl_pr.addnext(grid)
    for col in list(grid.findall(qn("w:gridCol"))):
        grid.remove(col)
    for w in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 567)))
        grid.append(col)


def set_table_cell_margins(table, top=70, start=90, bottom=70, end=90):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")




def set_row_height(row, height_dxa):
    """Set minimum row height in DXA units"""
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = OxmlElement("w:trHeight")
    tr_h.set(qn("w:val"), str(height_dxa))
    tr_pr.append(tr_h)

def apply_section_layout(section, compact=False):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    if compact:
        section.top_margin = Cm(0.6)
        section.bottom_margin = Cm(0.6)
        section.left_margin = Cm(1.15)
        section.right_margin = Cm(1.15)
    else:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)


def apply_run_font(run, size=16, bold=False):
    run.bold = bold
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)


def format_paragraph_runs(p, size=16, bold=False):
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        apply_run_font(run, size=size, bold=bold or bool(run.bold))


def cell_paragraph(cell, text="", size=12.5, bold=False, align=None):
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.text = text
    if align is not None:
        p.alignment = align
    format_paragraph_runs(p, size=size, bold=bold)
    return p


def paragraph(doc, text="", bold=False, align=None, size=16, before=0, after=0, first_line=None, line_spacing=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line_spacing == 1.5:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif line_spacing == 2.0:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        p.paragraph_format.line_spacing = line_spacing
    if first_line is not None:
        p.paragraph_format.first_line_indent = Cm(first_line)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    apply_run_font(run, size=size, bold=bold)
    return p


def add_runs(p, chunks, size=16):
    for text, bold in chunks:
        run = p.add_run(text)
        apply_run_font(run, size=size, bold=bold)


def setup_document():
    doc = Document()
    apply_section_layout(doc.sections[0], compact=True)
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(16)
    return doc


def add_page_break(doc, compact=False):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_section_layout(section, compact=compact)
    return section


def add_items_table(doc, items, include_details=False, min_rows=3, font_size=15, compact=False,
                    start_no=1, show_total=True, total_override=None):
    headers = ["ที่", "รายการและรายละเอียดคุณลักษณะเฉพาะพัสดุ" if include_details else "รายการ", "จำนวนหน่วย", "ราคา/หน่วย", "จำนวนเงิน", "หมายเหตุ"]
    if compact:
        widths = [0.9, 7.3 if include_details else 6.8, 1.6 if include_details else 1.8, 1.9 if include_details else 2.1, 2.4 if include_details else 2.5, 1.9]
    else:
        widths = [0.9, 7.3 if include_details else 6.8, 1.6 if include_details else 1.8, 1.9 if include_details else 2.1, 2.4 if include_details else 2.5, 1.9]
    rows = max(min_rows, len(items))
    extra_rows = 2 if show_total else 1  # header + (total row?)
    table = doc.add_table(rows=rows + extra_rows, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, sum(widths))
    set_table_fixed_layout(table, widths)
    # ลดระยะขอบบน-ล่างในเซลล์ ให้แถวเตี้ยลง (LibreOffice เรนเดอร์แถวสูงกว่า Word)
    # เพื่อให้ตาราง 15 แถว + ลายเซ็น พอดีในหน้าเดียว
    set_table_cell_margins(table, top=32, start=90, bottom=32, end=90)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, "D9EAF7")
        set_cell_width(cell, widths[i])
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                apply_run_font(run, size=font_size, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    total = 0
    for idx in range(rows):
        item = items[idx] if idx < len(items) else {}
        qty = money(item.get("quantity", 0))
        price = money(item.get("unitPrice", 0))
        amount = qty * price if item else 0
        total += amount
        values = [str(start_no + idx), clean_text(item.get("name", "")), clean_text(item.get("quantity", "")), fmt_money(price) if item else "", fmt_money(amount) if item else "", clean_text(item.get("note", ""))]
        for col, value in enumerate(values):
            cell = table.cell(idx + 1, col)
            set_cell_width(cell, widths[col])
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col in (1, 5) else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    apply_run_font(run, size=font_size)
    if show_total:
        total_row = table.rows[-1]
        total_row.cells[0].merge(total_row.cells[3])
        total_row.cells[0].text = "รวม"
        total_row.cells[4].text = fmt_money(total if total_override is None else total_override)
        total_row.cells[5].text = ""
        for i, cell in enumerate(total_row.cells):
            set_cell_shading(cell, "F2F2F2")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i < 4 else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    apply_run_font(run, size=font_size, bold=True)
    return total


def add_summary_table(doc, total_amount, font_size=14):
    """Summary table matching the example: ที่, รายการ, จำนวนหน่วย, ราคา/หน่วย, จำนวนเงิน, หมายเหตุ"""
    # Widths from example file (DXA->cm): 556, 4192, 1108, 1300, 1723, 1723
    widths = [0.981, 7.393, 1.954, 2.293, 3.039, 3.039]
    headers = ["ที่", "รายการ", "จำนวนหน่วย", "ราคา/หน่วย", "จำนวนเงิน", "หมายเหตุ"]

    table = doc.add_table(rows=3, cols=6)
    table.style = "Table Grid"
    set_table_width(table, sum(widths))
    set_table_fixed_layout(table, widths)
    set_table_cell_margins(table, top=75, start=90, bottom=75, end=90)

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "D9EAF7")
        set_cell_width(cell, widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_paragraph(cell, header, size=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data row
    data_values = ["1", "ดังเอกสารที่แนบมา", "1", fmt_money(total_amount), fmt_money(total_amount), ""]
    for col, value in enumerate(data_values):
        cell = table.cell(1, col)
        set_cell_width(cell, widths[col])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_paragraph(cell, value, size=font_size, align=WD_ALIGN_PARAGRAPH.LEFT if col == 1 else WD_ALIGN_PARAGRAPH.CENTER)

    # Total row - merge first 4 cells
    total_row = table.rows[2]
    total_row.cells[0].merge(total_row.cells[3])
    total_row.cells[0].text = ""
    for p in total_row.cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        apply_run_font(p.add_run("รวม"), size=font_size, bold=True)
    set_cell_shading(total_row.cells[0], "F2F2F2")
    set_cell_width(total_row.cells[0], sum(widths[:4]))

    cell_4 = total_row.cells[4]
    set_cell_shading(cell_4, "F2F2F2")
    set_cell_width(cell_4, widths[4])
    cell_paragraph(cell_4, fmt_money(total_amount), size=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    cell_5 = total_row.cells[5]
    set_cell_shading(cell_5, "F2F2F2")
    set_cell_width(cell_5, widths[5])
    cell_paragraph(cell_5, "", size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_signature(doc, name, title, label="ลงชื่อ........................................................", size=16, after=4, bind_next=False):
    p1 = paragraph(doc, label, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, size=size)
    p2 = paragraph(doc, f"( {name} )" if name else "(........................................................)", align=WD_ALIGN_PARAGRAPH.CENTER, after=0, size=size)
    p3 = paragraph(doc, f"ตำแหน่ง {title}" if title else "ตำแหน่ง........................................................", align=WD_ALIGN_PARAGRAPH.CENTER, after=after, size=size)
    # มัด 3 บรรทัดของลายเซ็นเดียวไว้ด้วยกัน ไม่ให้ตัดแยกคนละหน้า
    for p in (p1, p2, p3):
        p.paragraph_format.keep_together = True
    p1.paragraph_format.keep_with_next = True
    p2.paragraph_format.keep_with_next = True
    # bind_next=True เพื่อร้อยลายเซ็นนี้ให้อยู่หน้าเดียวกับลายเซ็น/หัวข้อถัดไป
    p3.paragraph_format.keep_with_next = bind_next
    return p3


def add_paginated_items(doc, items, sign_block, include_details=False, font_size=15,
                        per_page=15, money_text=None, heading=None):
    """แสดงตารางรายการพร้อมบล็อกลายเซ็นต่อท้าย
    - ถ้ารายการ <= per_page : ตาราง + ลายเซ็น อยู่หน้าเดียวกัน
    - ถ้ารายการ > per_page  : แบ่งเป็นหน้าๆ หน้าละ per_page รายการ และใส่ลายเซ็นทุกหน้า
      (ยอดรวมแสดงเฉพาะหน้าสุดท้าย) เพื่อให้กรรมการเห็นและเซ็นได้ทุกหน้า
    sign_block(doc) = ฟังก์ชันวาดบล็อกลายเซ็นของเอกสารหน้านั้น
    """
    grand_total = sum(money(i.get("quantity", 0)) * money(i.get("unitPrice", 0)) for i in items)
    if len(items) <= per_page:
        chunks = [items]
    else:
        chunks = [items[i:i + per_page] for i in range(0, len(items), per_page)]
    last = len(chunks) - 1
    start_no = 1
    for ci, chunk in enumerate(chunks):
        is_last = (ci == last)
        # หน้าเดียว: เติมแถวว่างให้สวยเหมือนเดิม / หลายหน้า: ใช้จำนวนแถวตามจริง
        min_rows = max(10, len(chunk)) if len(chunks) == 1 else len(chunk)
        add_items_table(doc, chunk, include_details=include_details, min_rows=min_rows,
                        font_size=font_size, start_no=start_no, show_total=is_last,
                        total_override=grand_total)
        start_no += len(chunk)
        if is_last and money_text:
            mp = paragraph(doc, money_text, after=8)
            mp.paragraph_format.keep_with_next = True
        if heading:
            hp = paragraph(doc, heading, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            hp.paragraph_format.keep_with_next = True
        sign_block(doc)
        if not is_last:
            add_page_break(doc)


def add_signature_to_cell(cell, name, title, label="ลงชื่อ........................................................", size=12.5):
    cell_paragraph(cell, label, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_paragraph(cell, f"( {name} )" if name else "(........................................................)", size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_paragraph(cell, f"ตำแหน่ง {title}" if title else "ตำแหน่ง........................................................", size=size, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_memo_header(doc, subject, doc_code, data, body_size=16, title_size=22, garuda_position=None):
    """Header with Garuda + บันทึกข้อความ"""
    if garuda_position == "left" and GARUDA_PATH.exists():
        # หน้า 1: ตราครุฑซ้าย + tabs + บันทึกข้อความ (ในบรรทัดเดียว)
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(0)
        p_title.paragraph_format.line_spacing = 1.0
        run_img = p_title.add_run()
        run_img.add_picture(str(GARUDA_PATH), width=Cm(2.02), height=Cm(2.25))
        run_tabs = p_title.add_run('\t\t\t\t\t')
        apply_run_font(run_tabs, size=title_size)
        run_title = p_title.add_run('บันทึกข้อความ')
        run_title.bold = True
        apply_run_font(run_title, size=title_size, bold=True)
    else:
        # หน้าอื่น: ตราครุฑซ้าย (แยก paragraph) + บันทึกข้อความ กึ่งกลาง
        if GARUDA_PATH.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after = Pt(0)
            p_img.paragraph_format.line_spacing = 1.0
            p_img.add_run().add_picture(str(GARUDA_PATH), width=Cm(2.02), height=Cm(2.25))
        paragraph(doc, 'บันทึกข้อความ', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=title_size, after=0, line_spacing=1.5)

    paragraph(doc, "ส่วนราชการ  โรงเรียนนายางกลักพิทยาคม  อำเภอเทพสถิต  จังหวัดชัยภูมิ", bold=True, size=body_size, after=0)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_runs(p, [
        (f"ที่ {doc_code}                            ", True),
        (f"วันที่ {data['day']} เดือน {data['month']} พ.ศ. {data['year']}", False),
    ], size=body_size)

    paragraph(doc, f"เรื่อง  {subject}", bold=True, size=body_size, after=0)
    paragraph(doc, "เรียน  ผู้อำนวยการโรงเรียนนายางกลักพิทยาคม", bold=True, after=0, size=body_size, line_spacing=1.5)


def add_first_page_review_sections(doc, data, total, remaining, font_size=16):
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_width(table, 18.7)
    set_table_cell_margins(table, top=60, start=100, bottom=60, end=100)
    # Row heights from example: 1485, 2971, 3344 DXA
    row_heights = [1485, 2971, 3344]
    for ri, row in enumerate(table.rows):
        set_row_height(row, row_heights[ri])
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, 9.35)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    left, right = table.rows[0].cells
    cell_paragraph(left, "ผู้รับผิดชอบ", size=font_size, bold=True)
    add_signature_to_cell(left, data["requesterName"], data["requesterPosition"], "ลงชื่อ........................................................ผู้รับผิดชอบ", size=font_size)
    cell_paragraph(right, "ผู้เห็นชอบ", size=font_size, bold=True)
    add_signature_to_cell(right, data["approverName"], f"หัวหน้ากลุ่มงานบริหาร{data['adminGroup']}", "ลงชื่อ........................................................ผู้เห็นชอบ", size=font_size)

    left, right = table.rows[1].cells
    cell_paragraph(left, "ได้ตรวจสอบแล้วมีในแผนจริง", size=font_size, bold=True)
    cell_paragraph(left, "1. ได้รับงบประมาณทั้งสิ้น    .............................. บาท", size=font_size)
    cell_paragraph(left, "2. ใช้จ่ายไปแล้ว              .............................. บาท", size=font_size)
    cell_paragraph(left, "3. ขอใช้ครั้งนี้              .............................. บาท", size=font_size)
    cell_paragraph(left, "4. คงเหลือ                   .............................. บาท", size=font_size)
    add_signature_to_cell(left, data["planOfficerName"], "ครู", "ลงชื่อ........................................................เจ้าหน้าที่แผนงาน", size=font_size)

    cell_paragraph(right, "ได้ตรวจสอบรายการ ขอซื้อ/ขอจ้าง แล้ว", size=font_size, bold=True)
    add_signature_to_cell(right, data["financeName"], "ครู", "ลงชื่อ........................................................เจ้าหน้าที่การเงิน", size=font_size)
    add_signature_to_cell(right, data["headOfficerName"], "ครู", "ลงชื่อ........................................................หัวหน้าเจ้าหน้าที่", size=font_size)

    left, right = table.rows[2].cells
    add_signature_to_cell(left, data["budgetAssistantName"], "หัวหน้ากลุ่มงานบริหารงบประมาณ", size=font_size)
    add_signature_to_cell(left, data["deputyName"], "รองผู้อำนวยการโรงเรียนนายางกลักพิทยาคม", size=font_size)

    cell_paragraph(right, "ความเห็นของผู้อำนวยการ", size=font_size, bold=True)
    cell_paragraph(right, "1. (    ) อนุมัติ      รายการที่........................", size=font_size)
    cell_paragraph(right, "2. (    ) ไม่อนุมัติ   รายการที่........................", size=font_size)
    cell_paragraph(right, "3. (    ) ทบทวนใหม่   รายการที่........................", size=font_size)
    cell_paragraph(right, "4. (    ) ดำเนินการตรวจสอบ รายการที่........................", size=font_size)
    add_signature_to_cell(right, data["directorName"], "ผู้อำนวยการโรงเรียนนายางกลักพิทยาคม", size=font_size)
    cell_paragraph(right, "วันที่............เดือน..............................พ.ศ..................", size=font_size)


def build_docx(data):
    items = [item for item in data.get("items", []) if clean_text(item.get("name"))]
    if not items:
        items = [{"name": "", "quantity": "", "unitPrice": "", "note": ""}]
    total = sum(money(i.get("quantity", 0)) * money(i.get("unitPrice", 0)) for i in items)
    data = {
        "day": clean_text(data.get("day"), str(datetime.now().day)),
        "month": clean_text(data.get("month"), THAI_MONTHS[datetime.now().month]),
        "year": clean_text(data.get("year"), str(datetime.now().year + 543)),
        "procurementType": clean_text(data.get("procurementType"), "ขอซื้อพัสดุ"),
        "requesterPrefix": clean_text(data.get("requesterPrefix"), "นาย"),
        "requesterName": clean_text(data.get("requesterName")),
        "requesterPosition": clean_text(data.get("requesterPosition"), "ครู"),
        "project": clean_text(data.get("project")),
        "adminGroup": clean_text(data.get("adminGroup")),
        "planPage": clean_text(data.get("planPage")),
        "totalBudget": fmt_money(data.get("totalBudget", total)) if money(data.get("totalBudget", 0)) != 0 else "............................",
        "spentBudget": fmt_money(data.get("spentBudget", 0)),
        "purpose": clean_text(data.get("purpose")),
        "deliveryDays": clean_text(data.get("deliveryDays"), "7"),
        "financeName": clean_text(data.get("financeName"), "นางสุจิตรา นามพิมล"),
        "planOfficerName": clean_text(data.get("planOfficerName"), "นายธีรพงษ์ พิมเคณา"),
        "approverName": clean_text(data.get("approverName")),
        "procurementOfficerName": clean_text(data.get("procurementOfficerName"), "นายอดิศักดิ์ วนาใส"),
        "headOfficerName": clean_text(data.get("headOfficerName"), "นายวิทวัช คำหา"),
        "budgetAssistantName": clean_text(data.get("budgetAssistantName"), "นางสุจิตรา นามพิมล"),
        "deputyName": clean_text(data.get("deputyName"), "นายธนกร วรรณชัย"),
        "directorName": clean_text(data.get("directorName"), "นางสาวสายฝน ทวีแก้ว"),
        "priceCommittee1Name": clean_text(data.get("priceCommittee1Name"), clean_text(data.get("committee1Name"))),
        "priceCommittee1Position": clean_text(data.get("priceCommittee1Position"), clean_text(data.get("committee1Position"), "ครู")),
        "priceCommittee2Name": clean_text(data.get("priceCommittee2Name"), clean_text(data.get("committee2Name"))),
        "priceCommittee2Position": clean_text(data.get("priceCommittee2Position"), clean_text(data.get("committee2Position"), "ครู")),
        "priceCommittee3Name": clean_text(data.get("priceCommittee3Name"), clean_text(data.get("committee3Name"))),
        "priceCommittee3Position": clean_text(data.get("priceCommittee3Position"), clean_text(data.get("committee3Position"), "ครู")),
        "inspectCommittee1Name": clean_text(data.get("inspectCommittee1Name")),
        "inspectCommittee1Position": clean_text(data.get("inspectCommittee1Position"), "ครู"),
        "inspectCommittee2Name": clean_text(data.get("inspectCommittee2Name")),
        "inspectCommittee2Position": clean_text(data.get("inspectCommittee2Position"), "ครู"),
        "inspectCommittee3Name": clean_text(data.get("inspectCommittee3Name")),
        "inspectCommittee3Position": clean_text(data.get("inspectCommittee3Position"), "ครู"),
    }
    remaining = money(data["totalBudget"]) - money(data["spentBudget"]) - total
    doc = setup_document()
    compact_size = 16
    compact_signature_size = 16

    # ===== หน้า 1 =====
    add_memo_header(doc, "ขอซื้อพัสดุ/ขอจ้างทำของ/ขอจ้างเหมาบริการ", "ศธ04299.37/จัดซื้อจัดจ้าง", data, body_size=compact_size, title_size=22, garuda_position="left")

    # P5: ด้วยข้าพเจ้า...สำหรับใช้ในโรงเรียน... (ทั้งหมดในพารากราฟเดียว เหมือนไฟล์ตัวอย่าง)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    add_runs(p, [
        ("ด้วยข้าพเจ้า ", False),
        (f"( {data['requesterPrefix']} ) {data['requesterName']} ", True),
        ("ตำแหน่ง ", False),
        (data["requesterPosition"], True),
        (" เป็นผู้รับผิดชอบ งาน/โครงการ ", False),
        (data["project"], True),
        (" กลุ่มบริหาร ", False),
        (data["adminGroup"], True),
        (" ซึ่งปรากฏในแผนปฏิบัติการของโรงเรียน หน้า ", False),
        (data["planPage"], True),
        (" ได้รับงบประมาณทั้งสิ้น ", False),
        (data["totalBudget"], True),
        (" บาท งบประมาณที่ขอใช้ในการดำเนินการครั้งนี้ ", False),
        (fmt_money(total), True),
        (" บาท ขอเสนอรายการประมาณการ เพื่อ ", False),
        (data["purpose"], True),
        (" สำหรับใช้ในโรงเรียนนายางกลักพิทยาคม โดยมีรายละเอียดดังนี้", False),
    ], size=compact_size)

    # Summary table (1 แถว แบบไฟล์ตัวอย่าง)
    add_summary_table(doc, total, font_size=12)

    paragraph(doc, f"จำนวนเงินตัวอักษร  ( {baht_text(total)} )", after=0, size=compact_size, line_spacing=1.5)

    paragraph(doc, "จึงเรียนมาเพื่อโปรดทราบและพิจารณา", first_line=1.25, size=compact_size, after=0, line_spacing=1.5)
    add_first_page_review_sections(doc, data, total, remaining, font_size=compact_signature_size)

    # ===== หน้า 2 =====
    add_page_break(doc)
    paragraph(doc, "เอกสารแนบท้ายบันทึกข้อความ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20)
    paragraph(doc, "ขอซื้อพัสดุ/ขอจ้างทำของ/ขอจ้างเหมาบริการ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    def sign_block_page2(d):
        add_signature(d, data["requesterName"], data["requesterPosition"], "ลงชื่อ........................................................ผู้รับผิดชอบ", bind_next=True)
        add_signature(d, data["approverName"], f"หัวหน้ากลุ่มงานบริหาร{data['adminGroup']}", "ลงชื่อ........................................................ผู้เห็นชอบ", bind_next=True)
        add_signature(d, data["directorName"], "ผู้อำนวยการโรงเรียนนายางกลักพิทยาคม")

    add_paginated_items(doc, items, sign_block_page2, money_text=f"จำนวนเงินตัวอักษร  {baht_text(total)}")

    # ===== หน้า 3 =====
    add_page_break(doc)
    add_memo_header(doc, "ขออนุมัติแต่งตั้งคณะกรรมการจัดทำราคากลางและคณะกรรมการตรวจรับพัสดุ", "ศธ04299.37/ราคากลางพัสดุ", data, title_size=22)
    paragraph(doc, f"ด้วย งาน/โครงการ {data['project']} มีความประสงค์จะดำเนินการ {data['procurementType']} ประจำปีงบประมาณ {data['year']} วงเงินงบประมาณ {fmt_money(total)} บาท", first_line=1.25)
    paragraph(doc, "ดังนั้น เพื่อให้การจัดทำราคากลางและรายละเอียดคุณลักษณะเฉพาะพัสดุ/กำหนดร่างขอบเขตของงาน เป็นไปตามพระราชบัญญัติการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 มาตรา 4 จึงเห็นสมควรแต่งตั้ง", first_line=1.25)
    paragraph(doc, "1. คณะกรรมการจัดทำราคากลางและรายละเอียดคุณลักษณะเฉพาะพัสดุ", bold=True)
    paragraph(doc, f"1.1 {data['priceCommittee1Name']} ตำแหน่ง {data['priceCommittee1Position']} ประธานกรรมการ")
    paragraph(doc, f"1.2 {data['priceCommittee2Name']} ตำแหน่ง {data['priceCommittee2Position']} กรรมการ")
    paragraph(doc, f"1.3 {data['priceCommittee3Name']} ตำแหน่ง {data['priceCommittee3Position']} กรรมการ")
    paragraph(doc, "มีหน้าที่จัดทำราคากลางและรายละเอียดคุณลักษณะเฉพาะพัสดุ/กำหนดร่างขอบเขตของงาน ที่จะ ซื้อ/จ้าง", first_line=1.25)
    paragraph(doc, "2. คณะกรรมการตรวจรับพัสดุ", bold=True)
    paragraph(doc, f"2.1 {data['inspectCommittee1Name']} ตำแหน่ง {data['inspectCommittee1Position']} ประธานกรรมการ/ผู้ตรวจรับพัสดุ  ลงชื่อ.................................")
    paragraph(doc, f"2.2 {data['inspectCommittee2Name']} ตำแหน่ง {data['inspectCommittee2Position']} กรรมการ  ลงชื่อ.................................")
    paragraph(doc, f"2.3 {data['inspectCommittee3Name']} ตำแหน่ง {data['inspectCommittee3Position']} กรรมการ  ลงชื่อ.................................")
    paragraph(doc, "มีหน้าที่ตรวจรับพัสดุตามใบสั่งซื้อ/ใบสั่งจ้าง หรือเอกสารอื่นที่เกี่ยวข้อง", first_line=1.25)
    closing_p3 = paragraph(doc, "จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ", first_line=1.25, line_spacing=1.5)
    closing_p3.paragraph_format.keep_with_next = True
    add_signature(doc, data["procurementOfficerName"], "เจ้าหน้าที่", "ลงชื่อ........................................................เจ้าหน้าที่", bind_next=True)
    add_signature(doc, data["headOfficerName"], "หัวหน้าเจ้าหน้าที่", "ลงชื่อ........................................................หัวหน้าเจ้าหน้าที่", bind_next=True)
    add_signature(doc, data["directorName"], "ผู้อำนวยการโรงเรียนนายางกลักพิทยาคม")

    # ===== หน้า 4 =====
    add_page_break(doc)
    add_memo_header(doc, f"ขอความเห็นชอบราคากลางและรายละเอียดคุณลักษณะเฉพาะพัสดุ โครงการ{data['project']}", "ศธ04299.37/ราคากลางพัสดุ", data, title_size=22)
    paragraph(doc, f"ตามบันทึกที่ ศธ04299.37/ราคากลางพัสดุ แต่งตั้งคณะกรรมการจัดทำราคากลางและรายละเอียดคุณลักษณะเฉพาะพัสดุ ในการ{data['procurementType']} เพื่อ{data['purpose']} นั้น โดยมีรายละเอียดดังนี้", first_line=1.25)
    paragraph(doc, f"จัดซื้อ/จ้าง ด้วยวิธีเฉพาะเจาะจง เนื่องจากมีวงเงินในการจัดซื้อจัดจ้างครั้งหนึ่งไม่เกินวงเงินตามที่กำหนด ราคากลางที่คำนวณได้ {fmt_money(total)} บาท วงเงินที่จะซื้อ/จ้าง {fmt_money(total)} บาท โดยพิจารณาคัดเลือกข้อเสนอโดยใช้เกณฑ์ราคา และผู้ขายจะต้องส่งมอบพัสดุภายในระยะเวลา {data['deliveryDays']} วัน", first_line=1.25)
    paragraph(doc, "บัดนี้ คณะกรรมการจัดทำราคากลาง ได้ดำเนินการจัดทำราคากลางและรายละเอียดคุณลักษณะเฉพาะพัสดุเรียบร้อยแล้ว ดังรายละเอียดที่แนบมาพร้อมนี้", first_line=1.25)
    closing_p4 = paragraph(doc, "จึงเรียนมาเพื่อโปรดพิจารณา", first_line=1.25, line_spacing=1.5)
    closing_p4.paragraph_format.keep_with_next = True
    add_signature(doc, data["priceCommittee1Name"], data["priceCommittee1Position"], "ลงชื่อ........................................................ประธานกรรมการ", bind_next=True)
    add_signature(doc, data["priceCommittee2Name"], data["priceCommittee2Position"], "ลงชื่อ........................................................กรรมการ", bind_next=True)
    add_signature(doc, data["priceCommittee3Name"], data["priceCommittee3Position"], "ลงชื่อ........................................................กรรมการ", bind_next=True)
    add_signature(doc, data["directorName"], "ผู้อำนวยการโรงเรียนนายางกลักพิทยาคม")

    # ===== หน้า 5 =====
    add_page_break(doc)
    paragraph(doc, "ราคากลางและรายละเอียดคุณลักษณะเฉพาะพัสดุ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20)
    paragraph(doc, f"งาน/โครงการ {data['project']} ของกลุ่มบริหาร {data['adminGroup']}")
    paragraph(doc, f"งบประมาณ {fmt_money(total)} บาท")

    def sign_block_page5(d):
        add_signature(d, data["priceCommittee1Name"], data["priceCommittee1Position"], "ลงชื่อ........................................................ประธานคณะกรรมการ", bind_next=True, size=15, after=2)
        add_signature(d, data["priceCommittee2Name"], data["priceCommittee2Position"], "ลงชื่อ........................................................กรรมการ", bind_next=True, size=15, after=2)
        add_signature(d, data["priceCommittee3Name"], data["priceCommittee3Position"], "ลงชื่อ........................................................กรรมการ", bind_next=True, size=15, after=2)
        add_signature(d, data["directorName"], "ผู้อำนวยการโรงเรียนนายางกลักพิทยาคม", size=15, after=2)

    add_paginated_items(doc, items, sign_block_page5, include_details=True,
                        money_text=f"จำนวนเงินตัวอักษร  {baht_text(total)}",
                        heading="คณะกรรมการกำหนดราคากลางและรายละเอียดคุณลักษณะเฉพาะพัสดุ")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()




def swap_font_for_pdf(docx_bytes, from_font="TH SarabunPSK", to_font="TH Sarabun New"):
    """Replace font name in the PDF-conversion copy only (the downloaded Word file is untouched).
    Maps the Word font 'TH SarabunPSK' to 'TH Sarabun New' which is the SIPA font bundled in
    fonts/ and installed in the Docker container, so LibreOffice renders Thai text identically."""
    import io, zipfile
    src = io.BytesIO(docx_bytes)
    dst = io.BytesIO()
    with zipfile.ZipFile(src, 'r') as zin, zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith('.xml') or item.filename.endswith('.rels'):
                try:
                    text = data.decode('utf-8')
                    text = text.replace(from_font, to_font)
                    data = text.encode('utf-8')
                except Exception:
                    pass
            zout.writestr(item, data)
    return dst.getvalue()


def build_pdf(data):
    """Generate DOCX then convert to PDF using Microsoft Word (docx2pdf) for pixel-perfect output"""
    import subprocess, tempfile
    docx_bytes = build_docx(data)
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        f.write(docx_bytes)
        tmp_docx = f.name
    tmp_pdf = tmp_docx.replace('.docx', '.pdf')

    # Try docx2pdf (uses Word COM on Windows — exact same output as Word "Save As PDF")
    try:
        from docx2pdf import convert
        convert(tmp_docx, tmp_pdf)
        if Path(tmp_pdf).exists():
            pdf_bytes = Path(tmp_pdf).read_bytes()
            Path(tmp_docx).unlink(missing_ok=True)
            Path(tmp_pdf).unlink(missing_ok=True)
            return pdf_bytes, True
    except Exception:
        pass

    # Fallback: LibreOffice (swap font for compatibility)
    pdf_docx_bytes = swap_font_for_pdf(docx_bytes)
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        f.write(pdf_docx_bytes)
        tmp_docx2 = f.name
    tmp_dir = Path(tmp_docx2).parent
    for exe in ['soffice', 'libreoffice',
                r'C:\Program Files\LibreOffice\program\soffice.exe',
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe']:
        try:
            result = subprocess.run(
                [exe, '--headless', '--convert-to', 'pdf', '--outdir', str(tmp_dir), tmp_docx2],
                capture_output=True, timeout=120,
            )
            tmp_pdf2 = tmp_docx2.replace('.docx', '.pdf')
            if result.returncode == 0 and Path(tmp_pdf2).exists():
                pdf_bytes = Path(tmp_pdf2).read_bytes()
                Path(tmp_docx2).unlink(missing_ok=True)
                Path(tmp_pdf2).unlink(missing_ok=True)
                Path(tmp_docx).unlink(missing_ok=True)
                return pdf_bytes, True
        except Exception:
            continue

    Path(tmp_docx).unlink(missing_ok=True)
    Path(tmp_docx2).unlink(missing_ok=True)
    return None, False

class Handler(BaseHTTPRequestHandler):
    def do_GET(self):
        parsed = urlparse(self.path)
        path = unquote(parsed.path.lstrip("/")) or "index.html"
        if path.startswith("generated/"):
            return self.serve_file(BASE_DIR / path, attachment=True)
        if path == "health":
            self.send_response(200)
            self.end_headers()
            self.wfile.write(b"ok")
            return
        target = BASE_DIR / path
        if not target.exists() or target.is_dir():
            target = BASE_DIR / "index.html"
        self.serve_file(target)

    def do_POST(self):
        if self.path == "/generate-pdf":
            return self.handle_generate_pdf()
        if self.path != "/generate":
            self.send_error(404)
            return
        length = int(self.headers.get("Content-Length", 0))
        payload = json.loads(self.rfile.read(length) or b"{}")
        docx_bytes = build_docx(payload)
        safe_project = re.sub(r"[^0-9A-Za-zก-๙._ -]+", "", clean_text(payload.get("project"), "เอกสารขอซื้อขอจ้าง")).strip()
        filename = f"{datetime.now().strftime('%Y%m%d-%H%M%S')}-{safe_project[:50] or 'procurement'}-{uuid.uuid4().hex[:6]}.docx"
        output_path = GENERATED_DIR / filename
        output_path.write_bytes(docx_bytes)
        body = json.dumps({"ok": True, "file": f"/generated/{filename}", "filename": filename}, ensure_ascii=False).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


    def handle_generate_pdf(self):
        length = int(self.headers.get("Content-Length", 0))
        payload = json.loads(self.rfile.read(length) or b"{}")
        # First generate DOCX
        docx_bytes = build_docx(payload)
        safe_project = re.sub(r"[^0-9A-Za-zก-๙._ -]+", "", clean_text(payload.get("project"), "procurement")).strip()
        base_name = f"{datetime.now().strftime('%Y%m%d-%H%M%S')}-{safe_project[:40] or 'procurement'}-{uuid.uuid4().hex[:6]}"
        docx_filename = base_name + ".docx"
        pdf_filename = base_name + ".pdf"
        docx_path = GENERATED_DIR / docx_filename
        docx_path.write_bytes(docx_bytes)
        # Convert to PDF
        pdf_bytes, ok = build_pdf(payload)
        if ok:
            pdf_path = GENERATED_DIR / pdf_filename
            pdf_path.write_bytes(pdf_bytes)
            body = json.dumps({"ok": True, "file": f"/generated/{pdf_filename}", "filename": pdf_filename, "docxFile": f"/generated/{docx_filename}"}, ensure_ascii=False).encode("utf-8")
        else:
            # Fallback: return docx info with error note
            body = json.dumps({"ok": False, "file": f"/generated/{docx_filename}", "filename": docx_filename, "error": "PDF conversion failed"}, ensure_ascii=False).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def serve_file(self, target, attachment=False):
        if not target.resolve().is_relative_to(BASE_DIR.resolve()) or not target.exists():
            self.send_error(404)
            return
        data = target.read_bytes()
        content_type = "text/html; charset=utf-8"
        if target.suffix == ".css":
            content_type = "text/css; charset=utf-8"
        elif target.suffix == ".js":
            content_type = "application/javascript; charset=utf-8"
        elif target.suffix == ".docx":
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif target.suffix == ".pdf":
            content_type = "application/pdf"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        if attachment:
            encoded_name = quote(target.name)
            fallback_name = "document" + target.suffix
            self.send_header("Content-Disposition", f"attachment; filename=\"{fallback_name}\"; filename*=UTF-8''{encoded_name}")
        self.end_headers()
        self.wfile.write(data)


def main():
    port = int(os.environ.get("PORT", "8087"))
    host = os.environ.get("HOST") or ("0.0.0.0" if os.environ.get("PORT") else "127.0.0.1")
    server = ThreadingHTTPServer((host, port), Handler)
    shown_host = "127.0.0.1" if host in ("", "0.0.0.0") else host
    print(f"ระบบพัสดุพร้อมใช้งาน: http://{shown_host}:{port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
